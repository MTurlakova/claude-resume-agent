"""
Reads resume content from resume_data.py and generates a formatted .docx file.
Run with: python generate_resume.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from resume_data import CONTACT_INFO, COMPANY, ROLE, OUTPUT_FOLDER, RESUME

# ── Formatting helpers ────────────────────────────────────────────────────────

def font(run, size=11, bold=False, italic=False):
    run.font.name = 'Garamond'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def sp(para, before=0, after=2):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)

def add_border(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def section_header(doc, title):
    p = doc.add_paragraph()
    r = p.add_run(title)
    font(r, bold=True)
    sp(p, before=6, after=2)
    add_border(p)

def bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Pt(-10)
    r = p.add_run(f'\u2022  {text}')
    font(r)
    sp(p, before=0, after=1)

# ── Build document ────────────────────────────────────────────────────────────

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(CONTACT_INFO['name'].upper())
font(r, size=16, bold=True)
sp(p, after=2)

# Contact line
contact_parts = [CONTACT_INFO.get('phone', ''), CONTACT_INFO.get('email', ''), CONTACT_INFO.get('location', '')]
contact_line = '  |  '.join(part for part in contact_parts if part)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(contact_line)
font(r)
sp(p, after=6)

# Summary
section_header(doc, 'SUMMARY')
p = doc.add_paragraph()
r = p.add_run(RESUME['summary'])
font(r)
sp(p, after=2)

# Experience
section_header(doc, 'EXPERIENCE')
for job in RESUME['experience']:
    p = doc.add_paragraph()
    r1 = p.add_run(job['employer'])
    font(r1, bold=True)
    location = f", {job['location']}" if job.get('location') else ''
    r2 = p.add_run(f"{location}  |  {job['dates']}")
    font(r2)
    sp(p, before=5, after=1)

    for i, role_entry in enumerate(job['roles']):
        p = doc.add_paragraph()
        r = p.add_run(role_entry['title'])
        font(r, italic=True)
        sp(p, before=0 if i == 0 else 3, after=1)
        for b in role_entry['bullets']:
            bullet(doc, b)

# Education
section_header(doc, 'EDUCATION')
for edu in RESUME['education']:
    p = doc.add_paragraph()
    r = p.add_run(f"{edu['institution']}, {edu['location']}")
    font(r, bold=True)
    sp(p, before=2, after=1)
    p = doc.add_paragraph()
    r = p.add_run(edu['degree'])
    font(r)
    sp(p, after=3)

# Skills
section_header(doc, 'SKILLS')
for skill in RESUME['skills']:
    p = doc.add_paragraph()
    r1 = p.add_run(skill['label'] + ' ')
    font(r1, bold=True)
    r2 = p.add_run(skill['content'])
    font(r2)
    sp(p, before=1, after=1)

# ── Save ──────────────────────────────────────────────────────────────────────

filename = f"{CONTACT_INFO['name']} Resume - {COMPANY} {ROLE}.docx"
filepath = os.path.join(OUTPUT_FOLDER, filename)
doc.save(filepath)
print(f"Saved to: {filepath}")
